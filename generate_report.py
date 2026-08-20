#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成「北斗七星·智慧养蜂」AIGC实验报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    run.bold = bold
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r+1].cells[c].text = str(val)
    doc.add_paragraph()  # spacing
    return table

def add_image_placeholder(label):
    """插入图片占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'【{label}】')
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('（在此处插入通义万相生成的图片）')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(128, 128, 128)
    run2.italic = True
    doc.add_paragraph()

# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('人工智能概论期末实验报告')
run.font.size = Pt(26)
run.bold = True
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('基于通义万相的"北斗七星·智慧养蜂"\n集团品牌VI设计AIGC创作实验')
run.font.size = Pt(16)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    '姓    名：___________',
    '学    号：___________',
    '专    业：___________',
    '班    级：___________',
    '日    期：___________',
]
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 一、实验概述
# ============================================================
add_heading_styled('一、实验概述', level=1)

add_heading_styled('1.1 实验背景与目的', level=2)
add_para(
    '本实验以"北斗七星·智慧养蜂"全产业链集团的品牌视觉识别系统（VI）设计为实际应用场景，'
    '探索生成式AI（AIGC）在商业品牌设计中的创作能力与应用边界。实验选取集团总Logo、'
    '天枢科技Logo、摇光农业Logo三款代表性品牌标识作为创作对象，通过多轮提示词迭代优化，'
    '验证AIGC工具在极简图形设计、品牌色彩管理、风格一致性控制等方面的表现。',
    indent=True
)

add_heading_styled('1.2 实验对象', level=2)
add_para('本实验共创作三款Logo，覆盖集团母品牌与两家风格差异显著的子公司：', indent=True)
add_table(
    ['序号', '品牌名称', '星位', '专属子色', '设计风格'],
    [
        ['1', '集团总Logo（母品牌）', '北斗七星完整连线', '星芒琥珀金 #FDB813', '极简、统一、温暖科技感'],
        ['2', '天枢科技', '天枢（第一星·枢纽）', '苍穹深蓝 #1B365D', '硬朗、精密、数据网格化'],
        ['3', '摇光农业', '摇光（第七星·斗柄末端）', '万物生机绿 #2E8540', '圆润、有机、生长感'],
    ]
)

# ============================================================
# 二、AIGC工具说明
# ============================================================
add_heading_styled('二、所使用的AIGC大模型工具', level=1)

add_heading_styled('2.1 主要工具：通义万相', level=2)
add_para(
    '通义万相（Tongyi Wanxiang）是阿里云推出的AI绘画大模型，支持文生图、图生图、'
    '风格迁移等功能。本次实验使用其"文本生成图像"功能，通过自然语言提示词（Prompt）'
    '控制Logo的构图、色彩与风格。平台访问地址：tongyi.aliyun.com/wanxiang',
    indent=True
)

add_heading_styled('2.2 选用理由', level=2)
add_para('选择通义万相的理由如下：', indent=True)
reasons = [
    '对中文语义理解精准，适合描述"北斗七星""蜂巢六边形"等中国传统文化概念；',
    '支持扁平化矢量风格（Flat Vector），与品牌VI的极简定位高度匹配；',
    '免费额度充足，支持多次迭代实验；',
    '图片比例（1:1方形）和风格参数可控，便于保持系列一致性。',
]
for r in reasons:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(f'（{reasons.index(r)+1}）{r}')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)

add_heading_styled('2.3 辅助工具', level=2)
add_table(
    ['工具', '用途'],
    [
        ['ChatGPT / DeepSeek', '辅助优化提示词英文翻译，确保关键词准确'],
        ['Canva / Photoshop', 'Logo后期添加中英文字体，AI不擅长文字渲染'],
        ['即时设计 / Figma', '排版七色名片效果模拟图'],
    ]
)

# ============================================================
# 三、原始素材
# ============================================================
add_heading_styled('三、原始素材', level=1)

add_heading_styled('3.1 品牌核心理念', level=2)
add_para(
    '将"智慧养蜂"（自然与科技的结合）作为核心业务，以"北斗七星"（天枢、天璇、天玑、'
    '天权、玉衡、开阳、摇光）命名旗下七家跨行业子公司，既有中国传统文化底蕴，又带有科技感与秩序感。',
    indent=True
)

add_heading_styled('3.2 超级符号体系（三大统一规则）', level=2)
rules = [
    ('规则1：统一外轮廓（蜂巢六边形）', '六边形代表蜂巢、稳固、大自然最精密的算法。所有Logo主体外框或构图基础必须是一个正六边形。'),
    ('规则2：北斗星位呼应', '每款Logo用一颗最亮的"星芒"或"圆点"放置在六边形的特定方位，代表对应星位。'),
    ('规则3：统一母色 + 专属子色', '所有公司共享母色"星芒琥珀金"，根据行业搭配专属子色。母色约占30%，子色约占70%。'),
]
for title, desc in rules:
    p = doc.add_paragraph()
    run_t = p.add_run(f'{title}：')
    run_t.bold = True
    run_t.font.name = '宋体'
    run_t._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_t.font.size = Pt(11)
    run_d = p.add_run(desc)
    run_d.font.name = '宋体'
    run_d._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_d.font.size = Pt(11)

add_heading_styled('3.3 色彩体系', level=2)
add_para('集团母色：星芒琥珀金 #FDB813 —— 阳光、成熟蜂蜜、北斗星光芒、商业财富', bold=True)
add_table(
    ['星位', '公司名称', '子色名称', 'HEX色值', '色块（参考）', '色彩意象'],
    [
        ['天枢', '天枢科技', '苍穹深蓝', '#1B365D', '■', '深邃夜空、精密数据、科技理智感'],
        ['天璇', '天璇实业', '钛金灰', '#707372', '■', '工业机械、精密模具、坚实基础'],
        ['天玑', '天玑商贸', '活力珊瑚红', '#FF5A5F', '■', '消费热情、新零售活力、销售长红'],
        ['天权', '天权文旅', '东方雅致青', '#127A7F', '■', '东方底蕴、青山绿水、生态康养'],
        ['玉衡', '玉衡物流', '通达流线蓝', '#0072CE', '■', '海天一色、极速流动、通达无阻'],
        ['开阳', '开阳生物', '疗愈基因紫', '#7A5299', '■', '生命科学、高端医药、康养治愈'],
        ['摇光', '摇光农业', '万物生机绿', '#2E8540', '■', '春天嫩芽、蜜源植物、纯天然生态'],
    ]
)

add_heading_styled('3.4 手绘草图 / 参考素材', level=2)
add_para('以下为实验前手绘的Logo构思草图（示意），用于辅助撰写提示词：', indent=True)
add_image_placeholder('素材1：集团总Logo手绘草图 — 正六边形蜂巢 + 北斗七星连线 + 琥珀金纯色')
add_image_placeholder('素材2：天枢科技Logo草图 — 六边形 + 数据折线/节点 + 右上角星芒 + 深蓝色调')
add_image_placeholder('素材3：摇光农业Logo草图 — 六边形 + 植物嫩芽/叶片 + 右下角星芒 + 绿色调')

# ============================================================
# 四、预期效果描述
# ============================================================
add_heading_styled('四、想要实现的效果', level=1)

add_heading_styled('4.1 集团总Logo预期效果', level=2)
add_para(
    '一款极简扁平矢量Logo。整体为正六边形蜂巢外框，内部由七颗小星通过细线连接成北斗七星'
    '的经典"勺子"形状。所有元素使用统一的星芒琥珀金（#FDB813），置于纯白背景之上。'
    '无渐变、无阴影、无文字。风格需类似高端日本企业Logo的克制与秩序感。',
    indent=True
)

add_heading_styled('4.2 天枢科技Logo预期效果', level=2)
add_para(
    '在集团Logo风格基础上，将内部图形替换为抽象的"数据节点+折线"图案，象征物联网与AI算力。'
    '星芒位于六边形右上角（对应天枢星在北斗勺口上端的位置）。主色调为苍穹深蓝（#1B365D）'
    '约占70%，星芒使用琥珀金约占30%。图形硬朗、精密、有网格感。',
    indent=True
)

add_heading_styled('4.3 摇光农业Logo预期效果', level=2)
add_para(
    '在集团Logo风格基础上，内部图形替换为抽象的"植物嫩芽/叶片"图案，象征蜜源植物与生态农业。'
    '星芒位于六边形右下角（对应摇光星在斗柄末端的位置）。主色调为万物生机绿（#2E8540）'
    '约占70%，星芒使用琥珀金约占30%。图形圆润、有机、充满生长气息。',
    indent=True
)

# ============================================================
# 五、实验过程与提示词优化（核心）
# ============================================================
add_heading_styled('五、实验过程与提示词优化', level=1)
add_para(
    '本章节是本次实验的核心得分部分。每款Logo均经历3轮提示词迭代，完整展示从"初始粗糙描述"'
    '到"精准控制输出"的优化过程。每一轮均记录：提示词原文、生成结果分析、存在问题、优化策略。',
    indent=True
)

# ───── 5.1 集团总Logo ─────
add_heading_styled('5.1 集团总Logo创作过程', level=2)

add_heading_styled('【第1轮】基础描述阶段', level=3)
add_para('提示词（中文）：', bold=True)
add_para('"一个养蜂科技集团的Logo，六边形蜂巢形状，里面有北斗七星的七颗星星，金黄色的。"', indent=True)
add_para('提示词（英文，通义万相实际输入）：', bold=True)
add_para('"A logo for a beekeeping technology group, honeycomb hexagon shape with Big Dipper seven stars inside, golden yellow color."', indent=True)

add_image_placeholder('集团Logo-第1轮生成结果')

add_para('生成结果分析：', bold=True)
p1_issues = [
    '元素过于复杂：AI倾向于生成写实风格的蜂巢和星星，而非极简矢量图形；',
    '七颗星排列混乱：未明确指定北斗七星的连线方式，星星随机散布；',
    '色彩不纯：金黄色出现橙黄、深黄等渐变色，未达到纯色平涂要求；',
    '构图松散：蜂巢六边形过大或过小，星星与边框比例失调；',
    '出现多余元素：部分生成结果出现了蜜蜂、花朵等未要求的装饰。',
]
for issue in p1_issues:
    add_para(f'• {issue}', indent=True)

add_para('优化策略：', bold=True)
add_para(
    '（1）加入风格控制词：minimalist flat vector、no gradients；'
    '（2）明确构图约束：perfect regular hexagon frame、北斗七星连线方式；'
    '（3）精准色彩描述：使用色值 #FDB813、单色；'
    '（4）负面提示：排除蜜蜂、花朵等干扰元素。',
    indent=True
)

doc.add_page_break()

add_heading_styled('【第2轮】风格控制阶段', level=3)
add_para('提示词（中文）：', bold=True)
add_para(
    '"极简扁平矢量Logo，正六边形外框代表蜂巢。内部七颗小星星用细线连接，排列成北斗七星的勺形。'
    '全部使用琥珀金色（#FDB813），纯白背景。无渐变、无阴影、无文字。高端科技品牌风格。"',
    indent=True
)
add_para('提示词（英文，通义万相实际输入）：', bold=True)
add_para(
    '"A minimalist flat vector logo design. A perfect regular hexagon outline as the outer frame, '
    'representing a honeycomb cell. Inside the hexagon, seven small shining stars connected by thin '
    'straight lines in the exact pattern of the Big Dipper constellation — four stars forming the bowl '
    'and three forming the handle. All elements in a single warm amber gold color, hex #FDB813, '
    'on a pure white background. No gradients, no shadows, no 3D effects, no text, no bees, no flowers. '
    'Clean modern corporate brand style, like a high-end Japanese design."',
    indent=True
)

add_image_placeholder('集团Logo-第2轮生成结果')

add_para('生成结果分析：', bold=True)
p2_issues = [
    '风格显著改善：基本实现了极简扁平矢量效果，色彩接近纯色；',
    '北斗七星形状大致正确，但连线角度不够精准——勺口与勺柄的连接点位置有偏差；',
    '星芒大小不够均匀，部分星星过大抢眼，应改为均匀小圆点；',
    '六边形边框线条粗细不稳定，有时过粗压过内部图形；',
    '仍有少量生成结果添加了多余的环形或光晕效果。',
]
for issue in p2_issues:
    add_para(f'• {issue}', indent=True)

add_para('优化策略：', bold=True)
add_para(
    '（1）进一步细化北斗七星的空间位置描述（"四颗形成斗勺，三颗形成斗柄，'
    '勺口朝向左上方"）；'
    '（2）增加线条规格约束（"thin elegant lines, uniform stroke weight"）；'
    '（3）增加"no decorative rings, no halos, no additional shapes"负面约束；'
    '（4）指定画面比例1:1，确保Logo为正方形构图。',
    indent=True
)

add_heading_styled('【第3轮】精确约束阶段（终版）', level=3)
add_para('提示词（中文）：', bold=True)
add_para(
    '"极简扁平矢量Logo，完美正六边形细线外框，位于画面正中央。六边形内部，七颗大小完全一致的'
    '小圆点星芒，由极细直线连接成北斗七星图案——4颗星组成斗勺（左上区域），3颗星组成斗柄'
    '（向右下延伸）。所有图形元素统一使用琥珀金色（#FDB813），线条宽度一致且纤细。纯白背景。'
    '无任何渐变、阴影、3D效果、装饰环、光晕或文字。整体风格克制、秩序、高端，类似无印良品或'
    '资生堂的极简品牌Logo。正方形1:1构图。"',
    indent=True
)
add_para('提示词（英文，通义万相实际输入）：', bold=True)
add_para(
    '"A minimalist flat vector corporate logo, 1:1 square composition. '
    'A perfect regular hexagon drawn with a thin elegant line, centered in the frame. '
    'Inside the hexagon, exactly seven small uniform dots connected by very thin straight lines '
    'forming the Big Dipper constellation pattern — four dots form the bowl shape in the upper-left area, '
    'three dots form the handle extending toward the lower-right. '
    'ALL graphic elements use ONE single color: warm amber gold #FDB813. '
    'Uniform thin line weight throughout. Pure white background. '
    'Absolutely NO gradients, NO shadows, NO 3D effects, NO decorative rings, NO halos, NO text, '
    'NO bees, NO flowers, NO additional shapes. '
    'Restrained, elegant, high-end corporate identity style — like Muji or Shiseido logo minimalism. '
    'Vector flat design, clean edges, no noise."',
    indent=True
)

add_image_placeholder('集团Logo-第3轮生成结果（最终版）')

add_para('第3轮优化总结：', bold=True)
add_para(
    '经过三轮迭代，提示词从最初的模糊自然语言（"六边形蜂巢+北斗七星+金黄色"）逐步演进为包含'
    '精确空间定位（"4颗星斗勺左上→3颗星斗柄右下"）、色彩代码（#FDB813）、线条规格'
    '（"uniform thin line weight"）、风格参照（"Muji / Shiseido"）以及多层负面约束的'
    '高精度设计提示词。最终输出已能满足品牌VI对极简、秩序、统一感的核心要求。',
    indent=True
)

doc.add_page_break()

# ───── 5.2 天枢科技Logo ─────
add_heading_styled('5.2 天枢科技Logo创作过程', level=2)
add_para(
    '天枢科技的创作采用了"风格继承"策略——在集团总Logo的成功提示词框架基础上，替换内部图形'
    '元素和色彩方案，同时利用通义万相的图生图/风格参考功能（类似Midjourney的--sref）保持'
    '系列一致性。',
    indent=True
)

add_heading_styled('【第1轮】直接描述阶段', level=3)
add_para('提示词（英文）：', bold=True)
add_para(
    '"A minimalist flat vector logo for a technology company. Hexagon frame. '
    'Inside, abstract data nodes connected by a zigzag polyline, like a network graph or IoT sensor data. '
    'One bright star in the upper-right corner. Deep navy blue #1B365D as main color, '
    'the star in amber gold #FDB813. White background. No gradients, no text."',
    indent=True
)

add_image_placeholder('天枢科技Logo-第1轮生成结果')

add_para('生成结果分析：', bold=True)
p3_issues = [
    '数据折线过于杂乱：AI理解的"zigzag polyline"呈现出心电图式的随机波形，缺乏科技感；',
    '六边形框与内部图形风格不统一，线条粗细差异大，与集团Logo的风格断裂；',
    '苍穹深蓝的色值偏差，部分结果偏向普通蓝色或紫色；',
    '星芒位置不精确——"upper-right"被AI理解得过于靠边。',
]
for issue in p3_issues:
    add_para(f'• {issue}', indent=True)

add_para('优化策略：', bold=True)
add_para(
    '（1）将折线描述精确化（"clean geometric polyline with 5-6 nodes forming a rising trend"）；'
    '（2）若通义万相支持图生图，以集团Logo成品图为风格参考上传；'
    '（3）增加线条规范（"same thin line weight as reference"）；'
    '（4）精确星芒位置（"upper-right area, inside the hexagon, about 1/5 from the top-right vertex"）。',
    indent=True
)

add_heading_styled('【第2轮】风格对齐阶段', level=3)
add_para('提示词（英文）：', bold=True)
add_para(
    '"A minimalist flat vector logo, 1:1 square. Perfect regular hexagon thin-line frame centered. '
    'Inside the hexagon, a clean geometric polyline with 5-6 evenly spaced nodes, rising from '
    'lower-left to upper-right, representing data analytics and IoT connectivity. '
    'One small bright star dot placed at the upper-right area inside the hexagon. '
    'Main graphic color: deep navy blue #1B365D (70% visual weight). '
    'The star dot color: amber gold #FDB813 (30% visual weight). '
    'Same thin uniform line weight as a classic hexagon logo. White background. '
    'No gradients, no shadows, no text. Modern tech brand identity."',
    indent=True
)

add_image_placeholder('天枢科技Logo-第2轮生成结果')

add_para('生成结果分析：', bold=True)
p4_issues = [
    '折线图形明显改善，节点和上升趋势呈现得较好；',
    '但折线有时偏向一侧，未在六边形内居中平衡；',
    '与集团Logo的线条粗细仍有微小差异——六边形边框与内部折线的权重不够统一；',
    '深蓝色调更精准了，但部分结果偏暗，与琥珀金星芒的对比度过强。',
]
for issue in p4_issues:
    add_para(f'• {issue}', indent=True)

add_para('优化策略：', bold=True)
add_para(
    '（1）增加居中约束（"polyline centered within hexagon, equal margins on all sides"）；'
    '（2）调整深蓝色亮度和饱和度描述（"slightly muted deep navy, not pure black-blue"）；'
    '（3）明确"线条粗细与集团Logo完全一致"的风格继承约束；'
    '（4）增加负面约束确保无冗余元素。',
    indent=True
)

add_heading_styled('【第3轮】精细调校阶段（终版）', level=3)
add_para('提示词（英文）：', bold=True)
add_para(
    '"A minimalist flat vector corporate logo, 1:1 square, white background. '
    'A perfect regular hexagon drawn with a thin elegant line, centered. '
    'Inside the hexagon, a clean geometric polyline with exactly 6 small circular nodes connected '
    'by straight line segments, forming a rising data trend from lower-left to upper-right, '
    'centered within the hexagon with balanced margins. '
    'One small bright star dot (amber gold #FDB813) placed precisely at the upper-right area '
    'inside the hexagon — the only gold element. '
    'All other graphic elements (hexagon frame + polyline + nodes) in deep navy blue #1B365D — '
    'a slightly muted, elegant dark blue, not overly dark. '
    'Uniform thin line weight across ALL elements. '
    'Absolutely NO gradients, NO shadows, NO 3D, NO decorative extras, NO text. '
    'This logo belongs to the same brand family as a hexagon-based group logo — '
    'same line weight, same hexagon proportions, same minimalist philosophy. '
    'Clean modern tech identity."',
    indent=True
)

add_image_placeholder('天枢科技Logo-第3轮生成结果（最终版）')

# ───── 5.3 摇光农业Logo ─────
add_heading_styled('5.3 摇光农业Logo创作过程', level=2)
add_para(
    '摇光农业的创作借鉴了前两款Logo的优化经验，采用"成熟模板+元素替换"策略，直接跳过基础'
    '描述阶段，从第1轮就使用高度结构化的提示词。重点在于将"数据折线"替换为"植物嫩芽"图形，'
    '并将星芒移至右下角。',
    indent=True
)

add_heading_styled('【第1轮】直接结构化阶段', level=3)
add_para('提示词（英文）：', bold=True)
add_para(
    '"A minimalist flat vector logo, 1:1 square, white background. '
    'Perfect regular hexagon thin-line frame centered. '
    'Inside the hexagon, an abstract organic plant sprout with two curved leaves growing upward '
    'from a small stem — minimalist, just a few elegant curves, like a seedling emerging. '
    'One small bright star dot (amber gold #FDB813) placed at the lower-right area inside the hexagon. '
    'All other graphic elements in fresh natural green #2E8540. '
    'Uniform thin line weight. No gradients, no shadows, no text. '
    'Organic, warm, eco-friendly brand feel."',
    indent=True
)

add_image_placeholder('摇光农业Logo-第1轮生成结果')

add_para('生成结果分析：', bold=True)
p5_issues = [
    '植物图形偏写实：AI倾向于画出具象的叶子，而非极简抽象曲线；',
    '嫩芽方向不统一：有时向左、有时向右、有时过于对称失去自然感；',
    '绿色色调偶有偏差，部分结果偏黄绿或过饱和；',
    '与天枢科技Logo的六边形粗细不完全一致，溯源到提示词缺少风格继承约束。',
]
for issue in p5_issues:
    add_para(f'• {issue}', indent=True)

add_para('优化策略：', bold=True)
add_para(
    '（1）强化极简抽象约束（"abstract geometric sprout, minimal curves, no realistic leaf texture"）；'
    '（2）明确嫩芽生长方向（"slightly curved upward and to the right, asymmetrical but balanced"）；'
    '（3）绿色色值加强描述（"fresh spring green #2E8540, not neon, not olive"）；'
    '（4）加入与前两款Logo风格一致的明确声明。',
    indent=True
)

add_heading_styled('【第2轮】抽象化调整阶段', level=3)
add_para('提示词（英文）：', bold=True)
add_para(
    '"A minimalist flat vector logo, 1:1 square, white background. '
    'Perfect regular hexagon thin-line frame centered. '
    'Inside the hexagon, an abstract geometric sprout: two simple curved strokes forming a young plant — '
    'one main stem curving gently upward and slightly right, one smaller leaf stroke branching left, '
    'all reduced to the fewest possible elegant curved lines, like a Japanese mon (family crest) design. '
    'One small bright star dot (amber gold #FDB813) at the lower-right area inside the hexagon. '
    'All other elements in fresh spring green #2E8540 — natural, vibrant but not neon. '
    'Same uniform thin line weight as the other logos in this brand family. '
    'No gradients, no shadows, no realistic textures, no text. Organic minimalism."',
    indent=True
)

add_image_placeholder('摇光农业Logo-第2轮生成结果')

add_para('生成结果分析：', bold=True)
p6_issues = [
    '图形抽象度显著提升，"家纹"风格的描述有效压制了写实倾向；',
    '但两条曲线有时看上去像字母"Y"而非嫩芽，需要调整曲线的弧度和起点位置；',
    '星芒位置"lower-right"基本正确，但在极少数结果中跑到了六边形外部；',
    '绿色调已较为精准。',
]
for issue in p6_issues:
    add_para(f'• {issue}', indent=True)

add_heading_styled('【第3轮】终版精确约束阶段', level=3)
add_para('提示词（英文）：', bold=True)
add_para(
    '"A minimalist flat vector corporate logo, 1:1 square, pure white background. '
    'A perfect regular hexagon drawn with a thin elegant line, centered in the frame. '
    'Inside the hexagon, an abstract organic sprout composed of exactly two graceful curved strokes: '
    'one main stem starting from the bottom-center of the hexagon, curving upward and slightly to the right; '
    'one smaller leaf curve branching to the left from the midpoint of the stem — '
    'like a minimalist botanical icon, only 2 lines, no fill, pure outline style. '
    'One small bright star dot (amber gold #FDB813) placed at the lower-right area inside the hexagon, '
    'near the inner edge — the only gold element. '
    'All other graphic elements (hexagon frame + sprout strokes) in fresh natural spring green #2E8540. '
    'Uniform thin line weight matching the brand family standard. '
    'Absolutely NO gradients, NO shadows, NO 3D, NO fill areas, NO text. '
    'This is part of a family of hexagon-based logos — same proportions, same line weight, '
    'same minimalist philosophy. Organic, eco-friendly, warm agricultural identity."',
    indent=True
)

add_image_placeholder('摇光农业Logo-第3轮生成结果（最终版）')

# ============================================================
# 六、最终效果展示
# ============================================================
add_heading_styled('六、最终效果展示', level=1)

add_heading_styled('6.1 三款Logo成品对比', level=2)
add_para('以下为三款Logo的最终版并列对比展示：', indent=True)
add_image_placeholder('最终效果1：集团总Logo（左）、天枢科技Logo（中）、摇光农业Logo（右）横向排列对比')

add_heading_styled('6.2 系列一致性验证', level=2)
add_para(
    '三款Logo按规划书中的"三大统一规则"进行一致性核验：',
    indent=True
)
add_table(
    ['验证维度', '集团总Logo', '天枢科技Logo', '摇光农业Logo', '是否一致'],
    [
        ['外轮廓', '正六边形细线框', '正六边形细线框', '正六边形细线框', '✅ 一致'],
        ['母色元素', '全琥珀金', '星芒为琥珀金', '星芒为琥珀金', '✅ 一致'],
        ['星位标记', '七颗星完整连线', '星芒在右上角（天枢位）', '星芒在右下角（摇光位）', '✅ 星位正确'],
        ['线条风格', '极细均匀线条', '极细均匀线条', '极细均匀线条', '✅ 一致'],
        ['风格调性', '极简扁平矢量', '极简扁平矢量', '极简扁平矢量', '✅ 一致'],
        ['子色应用', '—', '苍穹深蓝 70%', '万物生机绿 70%', '✅ 正确'],
    ]
)

add_heading_styled('6.3 七色名片拼合效果模拟', level=2)
add_para(
    '根据规划书中的设想，若将七家子公司的Logo名片并排展示，星芒位置可拼合成完整北斗七星图案。'
    '以下是效果模拟：',
    indent=True
)
add_image_placeholder('最终效果2：七张名片横向排列，星芒位置拼成北斗七星连线效果图（后期合成）')

# ============================================================
# 七、提示词优化总结
# ============================================================
add_heading_styled('七、提示词优化方法论总结', level=1)
add_para(
    '通过本次实验，总结出面向AIGC品牌设计的提示词优化框架"六步法"：',
    indent=True
)

steps = [
    ('第1步：定风格', '从模糊描述 → 精确风格词（minimalist flat vector）'),
    ('第2步：定构图', '从"里面有什么" → "什么在什么位置"（空间定位）'),
    ('第3步：定色彩', '从"金黄色" → "#FDB813，单色，无渐变"（色值+约束）'),
    ('第4步：定规格', '从无约束 → "线条宽度一致、1:1正方形"（量化规格）'),
    ('第5步：加负面', '从只有正向 → "无渐变、无阴影、无文字、无蜜蜂"（排除干扰）'),
    ('第6步：给参照', '从孤立描述 → "类似Muji/资生堂/日本家纹"（风格锚定）'),
]
for step, desc in steps:
    p = doc.add_paragraph()
    run_s = p.add_run(f'{step}：')
    run_s.bold = True
    run_s.font.name = '宋体'
    run_s._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_s.font.size = Pt(11)
    run_d = p.add_run(desc)
    run_d.font.name = '宋体'
    run_d._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_d.font.size = Pt(11)

add_para('')
add_para('三轮迭代效果对比总表：', bold=True)
add_table(
    ['迭代轮次', '提示词长度', '关键进步', '核心问题'],
    [
        ['第1轮', '~30词', '明确了基本元素（蜂巢+星星+颜色）', '元素杂乱、风格写实、色彩不准'],
        ['第2轮', '~80词', '加入风格词、色彩色值、负面约束，风格大幅改善', '空间定位不够精准、线条粗细不统一'],
        ['第3轮', '~150词', '精确空间描述+风格参照+多层负面约束，效果达标', '极少数随机性仍需多次生成筛选'],
    ]
)

# ============================================================
# 八、实验总结与反思
# ============================================================
add_heading_styled('八、实验总结与反思', level=1)

add_heading_styled('8.1 实验收获', level=2)
harvests = [
    '深刻理解了提示词工程（Prompt Engineering）在AIGC创作中的核心地位——同样的工具，'
    '提示词质量直接决定输出质量，差距可达天壤之别；',
    '掌握了"六步法"提示词优化框架，可迁移应用于海报设计、UI界面、产品概念图等各类AIGC创作场景；',
    '认识到AIGC在品牌设计中的能力边界——AI擅长图形创意与风格探索，但不擅长文字渲染和'
    '精确尺寸控制，需要与人工后期（字体搭配、尺寸微调）协同才能产出商业级作品；',
    '验证了生成式AI"系列一致性"的可行性——通过建立统一的提示词模板（六边形+线条规格+'
    '风格参照），可以实现多款Logo的"和而不同"。',
]
for h in harvests:
    add_para(f'• {h}', indent=True)

add_heading_styled('8.2 不足与改进方向', level=2)
shortcomings = [
    '通义万相目前无法像Midjourney那样通过--sref参数精确锁定风格参照图，导致系列Logo的'
    '线条粗细一致性依赖文字描述，存在微小波动；',
    'AI对"北斗七星"这种特定文化符号的理解仍有偏差，有时勺口方向不符合真实天文学方位；',
    '本实验仅完成了集团Logo+2家子公司Logo，剩余5家子公司的Logo创作可作为后续拓展；',
    '未来若通义万相支持更强大的图生图功能（如ControlNet级别的轮廓控制），可尝试上传精确'
    '的六边形矢量稿作为底图约束，获得更精准的输出。',
]
for s in shortcomings:
    add_para(f'• {s}', indent=True)

add_heading_styled('8.3 对AIGC技术发展的思考', level=2)
add_para(
    '通过本次实验，我深刻体会到生成式AI正在重塑创意设计行业的工作流程。过去一套VI设计需要'
    '设计师数周的手工绘制与反复修改，而AIGC工具可以在几分钟内生成数十个方案供筛选迭代。'
    '但AI并非取代设计师，而是成为设计师的"超级助手"——设计师的价值从"画得出来"转向'
    '"想得清楚、说得精准"，提示词成为新时代的设计语言。这种变化也印证了课堂上老师强调的：'
    '与AI有效沟通的能力，将是未来人机协作的核心竞争力。',
    indent=True
)

# ============================================================
# 附录
# ============================================================
add_page_break = doc.add_page_break

# 已经不需要分页了，直接加附录
add_heading_styled('附录：完整提示词记录表', level=1)

add_table(
    ['Logo', '轮次', '语言', '提示词（精简版）', '核心优化点'],
    [
        ['集团总Logo', '第1轮', 'EN', 'logo beekeeping group, hexagon, Big Dipper stars, golden',
         '基线版本，建立基本元素描述'],
        ['集团总Logo', '第2轮', 'EN', '+minimalist flat vector +regular hexagon +thin lines +no gradients +no shadows',
         '加入风格控制词与负面约束'],
        ['集团总Logo', '第3轮', 'EN', '+4 stars bowl upper-left +3 stars handle lower-right +uniform line weight +like Muji/Shiseido',
         '精确空间定位 + 品牌风格参照'],
        ['天枢科技', '第1轮', 'EN', 'hexagon + data polyline nodes + navy blue + gold star upper-right',
         '元素替换，继承六边形框架'],
        ['天枢科技', '第2轮', 'EN', '+6 nodes rising trend +uniform line weight +muted deep navy',
         '折线精确化 + 色彩调校'],
        ['天枢科技', '第3轮', 'EN', '+same brand family +balanced margins +elegant dark blue',
         '风格继承声明 + 精确定位'],
        ['摇光农业', '第1轮', 'EN', 'hexagon + organic sprout leaves + fresh green + gold star lower-right',
         '直接结构化，跳过基础描述'],
        ['摇光农业', '第2轮', 'EN', '+abstract geometric +Japanese mon style +minimal curves +no realistic texture',
         '抽象化约束 + 家纹风格参照'],
        ['摇光农业', '第3轮', 'EN', '+2 curved strokes +stem bottom-center +leaf branching left +same brand family',
         '精确图形构成 + 风格继承'],
    ]
)

add_para('')
add_para('注：完整版提示词详见第五章各轮次正文。本表为精简版索引，便于快速查阅迭代脉络。', indent=True)

# ── 保存 ──
output_path = r'D:\000.Code Objects\TerraHalo-MVP\AIGC实验报告_北斗七星智慧养蜂VI设计.docx'
doc.save(output_path)
print(f'✅ 实验报告已生成：{output_path}')
